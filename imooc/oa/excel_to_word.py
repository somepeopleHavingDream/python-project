"""解析excel，并将解析的内容转换成excel文件

"""


import xlrd
import random
# pip install --index-url https://pypi.douban.com/simple docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 1. 读取excel
data = xlrd.open_workbook('oa/data3.xlsx')
# 获取工作表
sheet = data.sheet_by_index(0)


class Question:
    pass


def create_question():
    """创建试题

    """

    question_list = []
    for i in range(sheet.nrows):
        if i > 1:
            obj = Question()
            # 题目
            obj.subject = sheet.cell(i, 1).value
            # 题型
            obj.question_type = sheet.cell(i, 2).value
            obj.option = []
            # a
            obj.option.append(sheet.cell(i, 3).value)
            # b
            obj.option.append(sheet.cell(i, 4).value)
            # c
            obj.option.append(sheet.cell(i, 5).value)
            # d
            obj.option.append(sheet.cell(i, 6).value)
            # 分值
            obj.score = sheet.cell(i, 7).value

            question_list.append(obj)
    
    # 将序列里所有元素随机排序
    random.shuffle(question_list)

    return question_list


def create_paper(file_name, paper_name, question_list):
    """生成word试卷

    """

    document = Document()

    # 页眉页脚的信息
    section = document.sections[0]
    header = section.header
    paragraph1 = header.paragraphs[0]
    paragraph1.text = paper_name
    footer = section.footer
    paragraph2 = footer.paragraphs[0]
    paragraph2.text = '内部试题，禁止泄露'

    # 试卷基本信息
    title = document.add_heading(paper_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph3 = document.add_paragraph()
    paragraph3.add_run('姓名：________')
    paragraph3.add_run('所属部门：________')
    paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 试题信息
    for question in question_list:
        subject = document.add_paragraph(style='List Number')
        run = subject.add_run(question.subject)
        # 加粗
        run.bold = True
        subject.add_run('【%s】分' % str(question.score))

        # 打乱选项的顺序
        random.shuffle(question.option)
        for index, option in enumerate(question.option):
            document.add_paragraph(('ABCD')[index] + str(option))
        document.save(file_name)


for i in range(2):
    question_list = create_question()
    create_paper('oa/paper' + str(i + 1) + '.docx', '2020年第一季度内部考试', question_list)